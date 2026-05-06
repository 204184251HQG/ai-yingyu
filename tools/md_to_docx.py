"""md_to_docx.py — 把项目下 docs/*.md 转成同名 .docx。

支持的 Markdown 语法：
- 标题 (# / ## / ### / ####)
- 段落（含 **bold**, *italic*, `code`, [link](url) 行内格式）
- 引用块 (> ...)
- 无序列表 (- / *) 和有序列表 (1. 2. 3.)
- 表格 (| col | col |)
- 代码块 (``` ... ```)
- 水平分隔线 (---)
- 任务/勾选框（☑/☐）按普通字符保留

用法：
    python tools/md_to_docx.py                # 转换 docs/*.md
    python tools/md_to_docx.py docs/foo.md    # 仅转换指定文件
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[错误] 缺少 python-docx, 请先运行: pip install python-docx")
    sys.exit(1)


IMAGE_LINE_RE = re.compile(r"^!\[([^\]]*)\]\(([^\)]+)\)\s*$")


HERE = Path(__file__).resolve().parent.parent
DOCS_DIR = HERE / "docs"

INLINE_PATTERN = re.compile(
    r"(\*\*[^*\n]+\*\*|__[^_\n]+__|\*[^*\n]+\*|_[^_\n]+_|`[^`\n]+`|\[[^\]]+\]\([^\)]+\))"
)


def _set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_runs_with_inline(paragraph, text, base_bold=False, base_italic=False):
    """把行内 markdown（粗/斜/代码/链接）转成多个 run。"""
    if not text:
        return
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.bold = base_bold
            run.italic = base_italic
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("__") and token.endswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.bold = base_bold
            run.italic = True
        elif token.startswith("_") and token.endswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.bold = base_bold
            run.italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith("["):
            inner = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token)
            if inner:
                label = inner.group(1)
                run = paragraph.add_run(label)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x99)
                run.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic


def _split_table_row(line):
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [c.strip() for c in raw.split("|")]


def _is_table_separator(line):
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-+:?", c.strip()) for c in cells)


def _add_table(doc, header, rows):
    n_cols = max(len(header), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    hdr = table.rows[0].cells
    for i in range(n_cols):
        cell = hdr[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_runs_with_inline(p, header[i] if i < len(header) else "", base_bold=True)
        _set_cell_shading(cell, "DCE6F1")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs_with_inline(p, row[c_idx] if c_idx < len(row) else "")
    doc.add_paragraph()


def _add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def _add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    _add_runs_with_inline(p, text, base_italic=True)
    run = p.runs[0] if p.runs else None
    if run:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def convert(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平线
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 图片：单独行 ![alt](path)，相对路径以 md 文件目录解析
        m = IMAGE_LINE_RE.match(stripped)
        if m:
            alt, src = m.group(1).strip(), m.group(2).strip()
            img_path = (md_path.parent / src).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                try:
                    run.add_picture(str(img_path), width=Inches(6.0))
                except Exception as e:
                    print(f"  [WARN] 插入图片失败 {img_path}: {e}")
                if alt:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(f"图：{alt}")
                    cr.italic = True
                    cr.font.size = Pt(10)
                    cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                print(f"  [WARN] 图片不存在: {img_path}")
                p = doc.add_paragraph(f"[缺失图片：{src}]")
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading(level=level)
            _add_runs_with_inline(heading, m.group(2).strip(), base_bold=True)
            i += 1
            continue

        # 引用块（连续 > 行合并）
        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                buf.append(lines[i].lstrip()[1:].lstrip())
                i += 1
            _add_quote(doc, " ".join(buf))
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _split_table_row(stripped)
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1
            _add_table(doc, header, rows)
            continue

        # 无序列表
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_inline(p, m.group(1).strip())
            i += 1
            continue

        # 有序列表
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_inline(p, m.group(1).strip())
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_runs_with_inline(p, stripped)
        i += 1

    if in_code and code_lines:
        _add_code_block(doc, code_lines)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  [OK] {md_path.name} -> {out_path.name}  ({out_path.stat().st_size} bytes)")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("inputs", nargs="*", help="md 文件路径，留空则转换 docs/*.md")
    ap.add_argument("--out-dir", default=str(DOCS_DIR), help="docx 输出目录")
    args = ap.parse_args()

    out_dir = Path(args.out_dir)
    if args.inputs:
        targets = [Path(p) for p in args.inputs]
    else:
        targets = sorted(DOCS_DIR.glob("*.md"))

    if not targets:
        print("[警告] 没有可转换的 md 文件")
        return

    for md_path in targets:
        out_path = out_dir / (md_path.stem + ".docx")
        convert(md_path, out_path)


if __name__ == "__main__":
    main()
